"""Tests for DOCXGenerator — markdown to Word document conversion."""

import pytest
from io import BytesIO
from docx import Document
from legal_service.application.services.docx_generator import DOCXGenerator


@pytest.fixture
def generator():
  return DOCXGenerator()


@pytest.fixture
def sample_markdown():
  return """# Mutual Non-Disclosure Agreement

**Date:** 10 March 2026

## 1. Purpose

The parties wish to explore potential business collaboration.

## 2. Duration

This Agreement shall remain in force for 2 years from the date hereof.

### 2.1 Early Termination

Either party may terminate by giving 30 days written notice.

## 3. Obligations

The Receiving Party shall:

- Keep all Confidential Information strictly confidential
- Not disclose to any third party without prior written consent
- Use the information only for the Purpose

## 4. Governing Law

This Agreement is governed by the laws of **England and Wales**.
"""


class TestCreateDocx:
  def test_returns_bytes_io(self, generator, sample_markdown):
    result = generator.create_docx(sample_markdown, 'Test NDA')
    assert isinstance(result, BytesIO)

  def test_produces_valid_docx(self, generator, sample_markdown):
    result = generator.create_docx(sample_markdown, 'Test NDA')
    doc = Document(result)
    assert doc is not None
    assert len(doc.paragraphs) > 0

  def test_contains_title(self, generator, sample_markdown):
    result = generator.create_docx(sample_markdown, 'Test NDA')
    doc = Document(result)
    texts = [p.text for p in doc.paragraphs]
    assert any('Mutual Non-Disclosure Agreement' in t for t in texts)

  def test_contains_headings(self, generator, sample_markdown):
    result = generator.create_docx(sample_markdown, 'Test NDA')
    doc = Document(result)
    texts = [p.text for p in doc.paragraphs]
    assert any('Purpose' in t for t in texts)
    assert any('Duration' in t for t in texts)

  def test_contains_body_text(self, generator, sample_markdown):
    result = generator.create_docx(sample_markdown, 'Test NDA')
    doc = Document(result)
    texts = [p.text for p in doc.paragraphs]
    assert any('potential business collaboration' in t for t in texts)

  def test_contains_list_items(self, generator, sample_markdown):
    result = generator.create_docx(sample_markdown, 'Test NDA')
    doc = Document(result)
    texts = [p.text for p in doc.paragraphs]
    assert any('Keep all Confidential Information' in t for t in texts)


class TestParseMarkdown:
  def test_parses_title(self, generator):
    sections = generator._parse_markdown('# My Title\n\nSome text')
    assert sections[0]['type'] == 'title'
    assert sections[0]['content'] == 'My Title'

  def test_parses_h2(self, generator):
    sections = generator._parse_markdown('## Section Two\n\nBody text')
    assert any(s['type'] == 'h2' for s in sections)

  def test_parses_h3(self, generator):
    sections = generator._parse_markdown('### Sub Section\n\nMore text')
    assert any(s['type'] == 'h3' for s in sections)

  def test_parses_list_items(self, generator):
    sections = generator._parse_markdown('- Item one\n- Item two')
    list_items = [s for s in sections if s['type'] == 'list_item']
    assert len(list_items) == 2

  def test_parses_paragraph(self, generator):
    sections = generator._parse_markdown('Just a paragraph of text.')
    assert any(s['type'] == 'paragraph' for s in sections)


class TestPageEstimate:
  def test_single_page_short_doc(self, generator):
    content = 'word ' * 200
    assert generator.estimate_page_count(content) == 1

  def test_multi_page_long_doc(self, generator):
    content = 'word ' * 600
    assert generator.estimate_page_count(content) >= 2

  def test_empty_content(self, generator):
    assert generator.estimate_page_count('') == 1


class TestEmptyContent:
  def test_empty_string_produces_valid_docx(self, generator):
    result = generator.create_docx('', 'Empty Doc')
    doc = Document(result)
    assert doc is not None
