"""DOCXGenerator — converts rendered markdown to professional Word documents.

Produces legal-formatted DOCX files with proper margins, fonts, and styling.
Uses python-docx for document creation.
"""

import re
import math
from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DOCXGenerator:
  """Converts markdown content to formatted Word documents."""

  FONT_NAME = 'Calibri'
  FONT_SIZE_BODY = Pt(11)
  FONT_SIZE_TITLE = Pt(16)
  FONT_SIZE_H2 = Pt(14)
  FONT_SIZE_H3 = Pt(12)
  LINE_SPACING = 1.15
  WORDS_PER_PAGE = 250

  def create_docx(self, markdown_content: str, title: str) -> BytesIO:
    """Convert markdown content to a DOCX document.

    Args:
      markdown_content: Rendered markdown string.
      title: Document title for metadata.

    Returns:
      BytesIO buffer containing the DOCX file.
    """
    doc = Document()

    # Set page margins
    for section in doc.sections:
      section.top_margin = Inches(1)
      section.bottom_margin = Inches(1)
      section.left_margin = Inches(1.25)
      section.right_margin = Inches(1.25)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = self.FONT_NAME
    font.size = self.FONT_SIZE_BODY

    # Parse and render sections
    if markdown_content:
      sections = self._parse_markdown(markdown_content)
      for section in sections:
        if section['type'] == 'title':
          self._add_title(doc, section['content'])
        elif section['type'] == 'h2':
          self._add_heading(doc, section['content'], level=2)
        elif section['type'] == 'h3':
          self._add_heading(doc, section['content'], level=3)
        elif section['type'] == 'list_item':
          self._add_list_item(doc, section['content'])
        elif section['type'] == 'paragraph':
          self._add_paragraph(doc, section['content'])

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

  def _parse_markdown(self, content: str) -> list[dict[str, str]]:
    """Parse markdown into typed sections.

    Returns:
      List of dicts with 'type' (title|h2|h3|list_item|paragraph)
      and 'content' keys.
    """
    sections: list[dict[str, str]] = []
    lines = content.split('\n')

    for line in lines:
      stripped = line.strip()
      if not stripped:
        continue

      if stripped.startswith('# ') and not stripped.startswith('## '):
        sections.append({'type': 'title', 'content': stripped[2:]})
      elif stripped.startswith('### '):
        sections.append({'type': 'h3', 'content': stripped[4:]})
      elif stripped.startswith('## '):
        sections.append({'type': 'h2', 'content': stripped[3:]})
      elif stripped.startswith('- ') or stripped.startswith('* '):
        sections.append({'type': 'list_item', 'content': stripped[2:]})
      else:
        sections.append({'type': 'paragraph', 'content': stripped})

    return sections

  def _add_title(self, doc: Document, text: str) -> None:
    """Add centered, bold title."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_after = Pt(12)
    run = para.add_run(text)
    run.bold = True
    run.font.size = self.FONT_SIZE_TITLE
    run.font.name = self.FONT_NAME

  def _add_heading(self, doc: Document, text: str, level: int = 2) -> None:
    """Add heading with level-specific formatting."""
    para = doc.add_paragraph()
    para.space_before = Pt(12)
    para.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.size = self.FONT_SIZE_H2 if level == 2 else self.FONT_SIZE_H3
    run.font.name = self.FONT_NAME

  def _add_paragraph(self, doc: Document, text: str) -> None:
    """Add body paragraph with inline markdown formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = self.LINE_SPACING
    self._add_formatted_text(para, text)

  def _add_list_item(self, doc: Document, text: str) -> None:
    """Add bullet list item."""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.line_spacing = self.LINE_SPACING
    self._add_formatted_text(para, text)

  def _add_formatted_text(self, para: Any, text: str) -> None:
    """Process inline markdown: **bold**, *italic*, "quoted definitions".

    Splits text on **bold** and *italic* markers and applies formatting
    to individual runs.
    """
    # Pattern: match **bold**, *italic*, or plain text segments
    pattern = r'(\*\*.*?\*\*|\*.*?\*|".*?")'
    parts = re.split(pattern, text)

    for part in parts:
      if not part:
        continue
      if part.startswith('**') and part.endswith('**'):
        run = para.add_run(part[2:-2])
        run.bold = True
        run.font.name = self.FONT_NAME
        run.font.size = self.FONT_SIZE_BODY
      elif part.startswith('*') and part.endswith('*'):
        run = para.add_run(part[1:-1])
        run.italic = True
        run.font.name = self.FONT_NAME
        run.font.size = self.FONT_SIZE_BODY
      elif part.startswith('"') and part.endswith('"'):
        run = para.add_run(part)
        run.bold = True
        run.font.name = self.FONT_NAME
        run.font.size = self.FONT_SIZE_BODY
      else:
        run = para.add_run(part)
        run.font.name = self.FONT_NAME
        run.font.size = self.FONT_SIZE_BODY

  def estimate_page_count(self, content: str) -> int:
    """Estimate page count based on word count (~250 words/page).

    Args:
      content: Document text content.

    Returns:
      Estimated page count (minimum 1).
    """
    word_count = len(content.split())
    if word_count == 0:
      return 1
    return max(1, math.ceil(word_count / self.WORDS_PER_PAGE))
